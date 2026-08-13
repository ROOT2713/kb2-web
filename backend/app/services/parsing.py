"""Document parsing service — PDF/Word/Excel/OCR pipeline.

Ported from: kb-web server.py parse_document() L815-L1034, mineru_parse_pdf() L654-L777,
             ocr_pdf() L590-L652, docx_to_pdf_via_libreoffice() L779-L809
"""

import asyncio
import logging
import os
import shutil
import subprocess
import tempfile
import traceback
import zipfile as _zipfile
from io import BytesIO
from pathlib import Path

import docx
import httpx
import pypdf

from app.config import settings

logger = logging.getLogger(__name__)

# ── MinerU 健康计数器 ──
mineru_stats = {"success": 0, "fail": 0, "last_error": None}
mineru_stats_lock = asyncio.Lock()


async def inc_mineru_success():
    async with mineru_stats_lock:
        mineru_stats["success"] += 1
        # 成功后重置失败计数，避免累计失败数导致永久degraded
        mineru_stats["fail"] = 0
        mineru_stats["last_error"] = None


async def inc_mineru_fail(err: str):
    async with mineru_stats_lock:
        mineru_stats["fail"] += 1
        mineru_stats["last_error"] = err[:200] if err else None


def get_mineru_stats() -> dict:
    """Thread-safe read of MinerU health stats. 用于 admin health endpoint."""
    return dict(mineru_stats)


def ocr_pdf(pdf_bytes: bytes) -> str:
    """用 pdftoppm + tesseract 对扫描件 PDF 做 OCR，返回纯文本"""
    tmpdir = tempfile.mkdtemp(prefix="kb_ocr_")
    try:
        pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)

        # 先探测页数
        page_count = len(pypdf.PdfReader(BytesIO(pdf_bytes)).pages)
        logger.info("OCR: %d 页扫描件，开始转换...", page_count)

        # PDF → 灰度 PNG（200 DPI 平衡速度与精度）
        try:
            result = subprocess.run(
                ["pdftoppm", "-png", "-gray", "-r", "200", pdf_path, os.path.join(tmpdir, "page")],
                capture_output=True, text=True, timeout=600,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"PDF 转换图片超时（{page_count} 页扫描件处理超过 5 分钟）。"
                "建议先用 PC 端工具将 PDF 压缩/降低分辨率后再上传。"
            )
        except Exception as e:
            logger.error("PDF 转换异常", exc_info=True)
            raise RuntimeError(f"PDF 转换异常: {e}")

        if result.returncode != 0:
            raise RuntimeError(f"pdftoppm 失败: {result.stderr[:200]}")

        # 逐页 OCR
        pages = sorted(Path(tmpdir).glob("page-*.png"))
        if not pages:
            raise RuntimeError("PDF 转换后无图片输出")

        texts = []
        for idx, png in enumerate(pages):
            if idx % 5 == 0 and idx > 0:
                logger.info("OCR: %d/%d 页已完成...", idx, len(pages))
            out_base = os.path.join(tmpdir, f"ocr_{png.stem}")
            try:
                ocr_result = subprocess.run(
                    ["tesseract", str(png), out_base, "-l", "chi_sim+eng", "--psm", "3"],
                    capture_output=True, text=True, timeout=60,
                )
            except subprocess.TimeoutExpired:
                logger.warning("OCR 页面 %s 超时，跳过", png.name)
                continue
            except Exception as e:
                logger.warning("OCR 页面 %s 异常: %s", png.name, e)
                continue
            if ocr_result.returncode != 0:
                logger.warning("OCR 页面 %s 失败: %s", png.name, ocr_result.stderr[:100])
                continue
            out_txt = out_base + ".txt"
            if os.path.exists(out_txt):
                with open(out_txt, "r", encoding="utf-8") as f:
                    texts.append(f.read().strip())
            os.unlink(out_txt)

        return "\n\n".join(texts)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


async def mineru_parse_pdf(filename: str, content: bytes) -> str:
    """通过 MinerU API 解析 PDF，返回 Markdown（含HTML表格）。

    流程：获取上传URL → PUT文件 → 轮询结果 → 下载ZIP → 提取 full.md

    超过 MINERU_PAGES_MAX 页自动分批，合并结果。
    """
    # 多 Key 并发分流：多批时 round-robin 绕开单 Key 排队瓶颈
    keys = [settings.mineru_api_key]
    if settings.mineru_api_key2:
        keys.append(settings.mineru_api_key2)
    if not keys[0]:
        raise RuntimeError("MINERU_API_TOKEN 未配置")

    reader = pypdf.PdfReader(BytesIO(content))
    total_pages = len(reader.pages)
    logger.info("MinerU: %d 页，%d 个Key可用，开始解析...", total_pages, len(keys))

    # 计算分批范围
    ranges = []
    for start in range(0, total_pages, settings.mineru_pages_max):
        if start >= total_pages:
            break
        end = min(start + settings.mineru_pages_max, total_pages)
        ranges.append((start + 1, end))  # page_ranges 是 1-indexed

    logger.info("MinerU: 分 %d 批: %s", len(ranges), ranges)

    all_md_parts = []
    failed_batches = []

    for batch_idx, (pg_start, pg_end) in enumerate(ranges):
        batch_label = f"batch{batch_idx+1}"
        if len(ranges) > 1:
            page_range = f"1-{pg_end - pg_start + 1}"
        else:
            page_range = f"{pg_start}-{pg_end}"

        mineru_base = settings.mineru_api_url or "https://mineru.net/api/v4"

        # Multi-Key retry: try each key, fallback on failure
        last_err = None
        for key_try in range(len(keys)):
            bearer = keys[(batch_idx + key_try) % len(keys)]
            if key_try > 0:
                logger.info("MinerU %s: Key%d 排队超时，切换到Key%d重试...",
                            batch_label,
                            (batch_idx + key_try - 1) % len(keys) + 1,
                            (batch_idx + key_try) % len(keys) + 1)
            try:
                # Step 1: 获取上传 URL
                async with httpx.AsyncClient(timeout=30) as client:
                    resp = await client.post(
                        f"{mineru_base}/file-urls/batch",
                        headers={
                            "Authorization": f"Bearer {bearer}",
                            "Content-Type": "application/json",
                        },
                        json={
                            "files": [{
                                "name": f"{batch_label}.pdf",
                                "is_ocr": True,
                                "page_ranges": page_range,
                            }],
                            "model_version": "pipeline",
                            "language": "ch",
                            "enable_table": True,
                        },
                    )
                    data = resp.json()
                    if data.get("code") != 0:
                        raise RuntimeError(f"MinerU 获取上传URL失败: {data.get('msg')}")

                    batch_id = data["data"]["batch_id"]
                    file_url = data["data"]["file_urls"][0]

                # Step 2: 上传文件
                if len(ranges) == 1:
                    upload_bytes = content
                else:
                    writer = pypdf.PdfWriter()
                    for pg in range(pg_start - 1, pg_end):
                        writer.add_page(reader.pages[pg])
                    upload_buf = BytesIO()
                    writer.write(upload_buf)
                    upload_bytes = upload_buf.getvalue()

                async with httpx.AsyncClient(timeout=120) as client:
                    resp = await client.put(file_url, content=upload_bytes)
                    if resp.status_code != 200:
                        raise RuntimeError(f"MinerU 文件上传失败: HTTP {resp.status_code}")

                logger.info("MinerU %s: 已上传 %d-%d 页", batch_label, pg_start, pg_end)

                # Step 3: 轮询结果
                poll_url = f"{mineru_base}/extract-results/batch/{batch_id}"
                max_wait = 900

                for _ in range(max_wait // 3):
                    await asyncio.sleep(3)
                    async with httpx.AsyncClient(timeout=30) as client:
                        resp = await client.get(
                            poll_url,
                            headers={"Authorization": f"Bearer {bearer}"},
                        )
                        data = resp.json()

                    if data.get("code") != 0:
                        continue

                    er = (data.get("data", {}).get("extract_result") or [{}])[0]
                    state = er.get("state")

                    if state == "done":
                        zip_url = er.get("full_zip_url")
                        break
                    elif state == "failed":
                        raise RuntimeError(f"MinerU {batch_label} 失败: {er.get('err_msg')}")
                else:
                    raise RuntimeError(f"MinerU {batch_label}: 轮询超时 ({max_wait}s)")

                # Step 4: 下载并解压 Markdown
                async with httpx.AsyncClient(timeout=60) as client:
                    resp = await client.get(zip_url)
                    zip_data = resp.content

                with _zipfile.ZipFile(BytesIO(zip_data)) as zf:
                    if "full.md" not in zf.namelist():
                        raise RuntimeError(f"MinerU {batch_label}: ZIP 中缺少 full.md")
                    md_text = zf.read("full.md").decode("utf-8")

                all_md_parts.append(md_text)
                logger.info("MinerU %s: Key%d 成功，%d 字符", batch_label,
                            (batch_idx + key_try) % len(keys) + 1, len(md_text))
                last_err = None
                break  # Success, exit retry loop

            except RuntimeError as e:
                last_err = e
                logger.warning("MinerU %s Key%d: %s", batch_label,
                               (batch_idx + key_try) % len(keys) + 1, e)
                continue

        if last_err:
            failed_batches.append(batch_idx)
            logger.error("MinerU %s: %d 个Key均失败", batch_label, len(keys))

    if failed_batches:
        if len(failed_batches) == len(ranges):
            raise RuntimeError(
                f"MinerU: 全部 {len(ranges)} 批失败（{len(keys)} 个Key均不可用），最近错误: {last_err}"
            )
        logger.warning("MinerU: %d/%d 批失败，使用部分结果", len(failed_batches), len(ranges))

    result = "\n\n".join(all_md_parts)
    logger.info("MinerU 完成: %d 字符", len(result))
    return result


async def docx_to_pdf_via_libreoffice(filename: str, content: bytes) -> bytes:
    """用 LibreOffice headless 将 DOCX 转为 PDF，返回 PDF 字节"""
    import tempfile
    tmpdir = tempfile.mkdtemp(prefix="kb_docx2pdf_")
    try:
        docx_path = os.path.join(tmpdir, filename)
        with open(docx_path, "wb") as f:
            f.write(content)

        proc = await asyncio.create_subprocess_exec(
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", tmpdir, docx_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)

        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice 退出码 {proc.returncode}: {stderr.decode()[:200]}")

        pdf_name = os.path.splitext(filename)[0] + ".pdf"
        pdf_path = os.path.join(tmpdir, pdf_name)
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"PDF 未生成: {pdf_path}")

        with open(pdf_path, "rb") as f:
            return f.read()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _odl_available() -> bool:
    try:
        import shutil
        return shutil.which("opendataloader-pdf") is not None
    except Exception:
        return False


async def opendataloader_parse_pdf(content: bytes, filename: str = "doc.pdf") -> str:
    """Use OpenDataLoader to parse PDF → structured Markdown. Requires Java 11+."""
    tmpdir = tempfile.mkdtemp(prefix="odl_")
    try:
        pdf_path = os.path.join(tmpdir, filename)
        with open(pdf_path, "wb") as f:
            f.write(content)
        out_dir = os.path.join(tmpdir, "out")
        os.makedirs(out_dir, exist_ok=True)
        logger.info("OpenDataLoader parsing: %s", filename)
        proc = await asyncio.create_subprocess_exec(
            "opendataloader-pdf", "-o", out_dir, "-f", "markdown", pdf_path,
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
        except asyncio.TimeoutError:
            proc.kill()
            raise RuntimeError("OpenDataLoader timed out (120s)")
        if proc.returncode != 0:
            err_msg = stderr.decode("utf-8", errors="replace")[:300]
            raise RuntimeError(f"OpenDataLoader exit {proc.returncode}: {err_msg}")
        md_files = sorted(Path(out_dir).glob("*.md"))
        if not md_files:
            raise RuntimeError("OpenDataLoader produced no Markdown output")
        text = md_files[0].read_text(encoding="utf-8")
        if not text.strip():
            raise RuntimeError("OpenDataLoader returned empty text")
        logger.info("OpenDataLoader done: %d chars", len(text))
        return text
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


async def parse_document(filename: str, content: bytes) -> str:
    """解析 PDF/Word/Markdown/TXT → 纯文本

    PDF 优先使用 MinerU API（高精度表格识别），失败时回退到 tesseract OCR。
    OpenDataLoader 作为 MinerU 后的第二 fallback（比 pypdf 保留更多表格/标题结构）。
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        try:
            reader = pypdf.PdfReader(BytesIO(content))
        except Exception as e:
            raise ValueError(f"PDF 解析失败（文件可能已损坏或加密）: {e}")
        if reader.is_encrypted:
            raise ValueError("PDF 文件已加密，无法提取文字内容")

        page_count = len(reader.pages)

        # ── PDF统一MinerU优先解析（pypdf仅作兜底）──
        if settings.mineru_api_key:
            try:
                logger.info("PDF MinerU 解析中 (%d 页)...", page_count)
                text = await mineru_parse_pdf(filename, content)
                if text and text.strip():
                    logger.info("MinerU 完成，提取 %d 字符", len(text))
                    await inc_mineru_success()
                    return text
                logger.info("MinerU 返回空结果，回退 pypdf")
            except Exception as e:
                logger.warning("MinerU 失败，回退 pypdf: %s", e)
                logger.error("", exc_info=True)
                await inc_mineru_fail(str(e))

        # ── Fallback: OpenDataLoader（数字PDF结构保留，比pypdf表格/标题完整）──
        if _odl_available():
            try:
                text = await opendataloader_parse_pdf(content, filename)
                if text and text.strip():
                    return text
            except Exception as e:
                logger.warning("OpenDataLoader 失败，回退 pypdf: %s", e)

        # ── 兜底: pypdf 提取文字层 ──
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
        if text.strip():
            logger.info("pypdf 提取 %d 字符", len(text))
            return text

        # ── 最后兜底: tesseract OCR ──
        logger.info("pypdf 文字层为空 (%d 页)，尝试 tesseract OCR...", page_count)
        try:
            text = ocr_pdf(content)
            if not text.strip():
                raise ValueError(
                    "PDF OCR 识别结果为空。"
                    "可能原因：①图片质量过低 ②PDF 为纯图片且文字不清晰。"
                    "建议检查 MinerU 服务是否正常。"
                )
            logger.info("tesseract 完成，提取 %d 字符", len(text))
        except RuntimeError as e:
            raise ValueError(f"PDF OCR 失败: {e}")
        except Exception as e:
            logger.error("PDF OCR 异常", exc_info=True)
            raise ValueError(f"PDF OCR 异常: {e}")
        return text
    elif ext in (".docx", ".doc"):
        # ── 优先: LibreOffice → PDF → MinerU（保留表格结构+版面分析）──
        try:
            logger.info("DOCX → PDF 转换中: %s", filename)
            pdf_bytes = await docx_to_pdf_via_libreoffice(filename, content)
            if pdf_bytes:
                pdf_name = filename.rsplit(".", 1)[0] + ".pdf"
                # 优先 MinerU 高精度解析
                if settings.mineru_api_key:
                    try:
                        text = await mineru_parse_pdf(pdf_name, pdf_bytes)
                        if text.strip():
                            logger.info("DOCX→PDF→MinerU 完成: %d 字符", len(text))
                            return text
                    except Exception as e:
                        logger.warning("MinerU 解析失败，回退 pypdf: %s", e)
                # MinerU 不可用或失败 → OpenDataLoader（结构保留好）
                if _odl_available() and pdf_bytes:
                    try:
                        text = await opendataloader_parse_pdf(pdf_bytes, pdf_name)
                        if text.strip():
                            logger.info("DOCX→PDF→OpenDataLoader 完成: %d 字符", len(text))
                            return text
                    except Exception as e:
                        logger.warning("OpenDataLoader 失败，回退 pypdf: %s", e)
                # → pypdf 提取文字层
                try:
                    reader = pypdf.PdfReader(BytesIO(pdf_bytes))
                    text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
                    if text.strip():
                        logger.info("DOCX→PDF→pypdf 完成: %d 字符", len(text))
                        return text
                except Exception as e:
                    logger.warning("pypdf 解析失败: %s", e)
        except Exception as e:
            logger.warning("DOCX→PDF 转换失败，回退 python-docx: %s", e)
        # ── Fallback: python-docx 直接解析（表格结构会丢失）──
        try:
            d = docx.Document(BytesIO(content))
        except Exception as e:
            raise ValueError(f"Word 文档解析失败: {e}")
        text = "\n\n".join(p.text for p in d.paragraphs)
        for table in d.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    text += "\n" + row_text
        if text.strip():
            logger.info("python-docx fallback 完成: %d 字符", len(text))
            return text
        raise ValueError("Word 文档内容为空。请检查文件是否包含可读文字。")
    elif ext in (".xlsx",):
        import openpyxl
        try:
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        except Exception as e:
            raise ValueError(f"Excel 文件解析失败: {e}")

        def _build_merged_map(ws):
            """Build a lookup dict: (row, col) -> top-left cell value for all merged ranges."""
            merged_map = {}
            for merged_range in ws.merged_cells.ranges:
                top_val = ws.cell(merged_range.min_row, merged_range.min_col).value
                for r in range(merged_range.min_row, merged_range.max_row + 1):
                    for c in range(merged_range.min_col, merged_range.max_col + 1):
                        merged_map[(r, c)] = top_val
            return merged_map

        def _cell_value(ws, row_idx, col_idx, merged_map):
            """Get cell value, handling merged cells and cleaning."""
            val = merged_map.get((row_idx, col_idx))
            if val is None:
                val = ws.cell(row_idx, col_idx).value
            if val is None:
                return ""
            s = str(val).strip()
            s = s.replace("\r\n", "；").replace("\n", "；").replace("\r", "；")
            return s

        all_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            merged_map = _build_merged_map(ws)
            sheet_lines = [f"[Sheet: {sheet_name}]"]

            # Find headers: first non-empty row, then merge sub-header rows
            headers = []
            header_rows_end = 0
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            for r in range(1, max_row + 1):
                row_vals = [_cell_value(ws, r, c, merged_map) for c in range(1, max_col + 1)]
                if not any(row_vals):
                    continue
                if not headers:
                    # First non-empty row = primary headers
                    headers = [v if v else "" for v in row_vals]
                    header_rows_end = r
                else:
                    # Check if this is a sub-header row (short values, no pure numbers)
                    has_long_text = any(len(v) > 15 for v in row_vals if v)
                    has_only_numbers = all(
                        v.replace(".", "").replace("-", "").isdigit() if v else True
                        for v in row_vals
                    )
                    if has_long_text or has_only_numbers:
                        break  # This is data, not a sub-header
                    # Merge sub-header into headers (e.g., "评分要点" + "要点分值" → "评分要点-要点分值")
                    # Skip merge if values are identical (avoid "序号-序号")
                    for ci, v in enumerate(row_vals):
                        if v and ci < len(headers):
                            if headers[ci] and headers[ci] != v:
                                headers[ci] = f"{headers[ci]}-{v}"
                            elif not headers[ci]:
                                headers[ci] = v
                    header_rows_end = r

            if not headers or not any(headers):
                if len(sheet_lines) > 0:
                    all_text.append("\n".join(sheet_lines))
                continue

            # Find special column indices
            xuhao_col = None  # 序号
            category_col = None  # 检查类别
            for ci, h in enumerate(headers):
                h_stripped = h.strip()
                if "序号" in h_stripped:
                    xuhao_col = ci
                elif "检查类别" in h_stripped:
                    category_col = ci

            # Process data rows (skip all header rows)
            for r in range(header_rows_end + 1, max_row + 1):
                # Check if row is completely empty
                row_values = [_cell_value(ws, r, c, merged_map) for c in range(1, max_col + 1)]
                if not any(row_values):
                    continue

                # Build item header line
                item_header_parts = []
                if xuhao_col is not None and row_values[xuhao_col]:
                    item_header_parts.append(f"第{row_values[xuhao_col]}项")
                if category_col is not None and row_values[category_col]:
                    item_header_parts.append(row_values[category_col])

                if item_header_parts:
                    sheet_lines.append(" - ".join(item_header_parts))

                # Output each non-empty cell as "列名: 值"
                for ci, val in enumerate(row_values):
                    if not val:
                        continue
                    col_name = headers[ci].strip()
                    # Skip 序号 and 检查类别 if already in header
                    if ci == xuhao_col or ci == category_col:
                        continue
                    sheet_lines.append(f"{col_name}: {val}")

                sheet_lines.append("")  # blank line between items

            # Only add sheet if there's more than just the title
            if len(sheet_lines) > 1:
                all_text.append("\n".join(sheet_lines))

        num_sheets = len(wb.sheetnames)
        wb.close()
        text = "\n\n".join(all_text)
        if not text.strip():
            raise ValueError("Excel 文件内容为空。请检查文件是否包含可读数据。")
        logger.info("openpyxl 提取 %d 字符 (%d sheets)", len(text), num_sheets)
        return text
    elif ext in (".md", ".markdown", ".txt", ".text", ""):
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
